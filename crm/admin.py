from .models import *
from .forms import *
from .views import financial_report
from django import forms
from django.contrib import admin, messages
from django.core.exceptions import ValidationError
from django.utils.html import format_html
from django.urls import path, reverse
from django.db.models import Count, Sum
import openpyxl
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from django.shortcuts import render, redirect
from openpyxl.styles import Font, Alignment, Border, Side
from urllib.parse import quote
from datetime import datetime


@admin.register(ServiceReport)
class ServiceReportAdmin(admin.ModelAdmin):
    change_list_template = 'admin/service_report_change_list.html'  # Шаблон для списка изменений
    date_hierarchy = 'appointment_date'  # Иерархия дат
    list_filter = ['personal__full_name']  # Фильтры списка

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        if request.GET:
            filters = {key: value for key, value in request.GET.items() if key in self.list_filter}
            qs = qs.filter(**filters)
        return qs.filter(status='выполнено')  # Фильтруем только выполненные записи

    def changelist_view(self, request, extra_context=None):
        response = super().changelist_view(request, extra_context=extra_context)
        try:
            qs = response.context_data['cl'].queryset
        except (AttributeError, KeyError):
            return response

        services = list(Appointment.objects.values_list('service__title', flat=True).distinct())  # Список уникальных услуг
        metrics = {'total_earnings': Sum('service__price')}  # Метрики для агрегации данных

        summary = {service: dict(
            qs.filter(service__title=service).values_list('personal__full_name').annotate(total_services=Count('id')).order_by('personal__full_name')
        ) for service in services}  # Резюме по услугам и сотрудникам

        earnings = qs.values('personal__full_name').annotate(**metrics).order_by('personal__full_name')  # Заработок по сотрудникам

        # Вычисляем данные для экспорта - делаем это только один раз!
        services_all, earnings_by_employee_all, total_services_all, total_earnings_all = self.process_data(qs)

        extra_context = {
            'summary': summary,
            'services': services,
            'earnings': earnings,
            'cl': response.context_data['cl'],
            'services_all': services_all,  # Передаем в функции экспорта
            'earnings_by_employee_all': earnings_by_employee_all,  # Передаем в функции экспорта
            'total_services_all': total_services_all,  # Передаем в функции экспорта
            'total_earnings_all': total_earnings_all,  # Передаем в функции экспорта
        }

        return super().changelist_view(request, extra_context=extra_context)

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('financial_report/', self.admin_site.admin_view(self.financial_report_view), name='financial_report'),
            path('export_excel/', self.admin_site.admin_view(self.date_filter_view), {'export_to': 'excel'}, name='export_to_excel'),
            path('export_docx/', self.admin_site.admin_view(self.date_filter_view), {'export_to': 'docx'}, name='export_to_docx'),
            path('date_filter/', self.admin_site.admin_view(self.date_filter_view), name='date_filter')
        ]
        return custom_urls + urls

    def has_add_permission(self, request):
        return False  # Запрещаем добавление записей

    def financial_report_view(self, request):
        pass  # Вью для финансового отчета (пустая пока)

    def apply_date_filters(self, queryset, start_date, end_date):
        if start_date:
            queryset = queryset.filter(appointment_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(appointment_date__lte=end_date)
        return queryset  # Применяем фильтры по дате

    def process_data(self, queryset, start_date=None, end_date=None): #обработка данных
        queryset = self.apply_date_filters(queryset, start_date, end_date)
        services = list(queryset.values_list('service__title', flat=True).distinct())  # Список уникальных услуг
        earnings_by_employee = {}  # Заработок по сотрудникам
        total_services = {service: 0 for service in services}  # Количество услуг по типам
        total_earnings = 0  # Общий заработок

        for item in queryset:
            employee_name = item.personal.full_name
            service_name = item.service.title
            price = item.service.price

            earnings_by_employee[employee_name] = earnings_by_employee.get(employee_name, 0) + price
            total_services[service_name] += 1
            total_earnings += price

        return services, earnings_by_employee, total_services, total_earnings


    def date_filter_view(self, request, export_to=None):
        if request.method == 'POST':
            form = DateRangeForm(request.POST)
            if form.is_valid():
                start_date = form.cleaned_data['start_date']
                end_date = form.cleaned_data['end_date']
                try:
                    if request.POST.get('export_type') == 'excel':
                        return self.export_to_excel(request, start_date, end_date)
                    elif request.POST.get('export_type') == 'docx':
                        return self.export_to_docx(request, start_date, end_date)
                except AttributeError as e:
                    messages.error(request, f"Ошибка: {e}. Проверьте корректность введенных дат.")
                    return render(request, 'admin/date_filter_form.html', {'form': form}) 
            else:
                # Отображаем ошибки формы
                return render(request, 'admin/date_filter_form.html', {'form': form})
        else:
            form = DateRangeForm()
        context = {'form': form}
        return render(request, 'admin/date_filter_form.html', context)  # Форма для выбора диапазона дат

    def export_to_excel(self, request, start_date, end_date):
        if not start_date or not end_date:
            messages.error(request, "Ошибка: Даты не выбраны.")
            return redirect('admin:index')
        queryset = self.get_queryset(request)
        services, earnings_by_employee, total_services, total_earnings = self.process_data(queryset, start_date, end_date)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Отчет по услугам"
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE  # Альбомная ориентация

        header_text = f"Отчет по услугам за период с {start_date.strftime('%d-%m-%Y')} по {end_date.strftime('%d-%m-%Y')}"
        ws.merge_cells('A1:C1')
        header_cell = ws.cell(row=1, column=1, value=header_text)
        header_cell.font = Font(name='Times New Roman', size=14, bold=True)
        header_cell.alignment = Alignment(horizontal='left')

        headers = ["Сотрудник"] + services + ["Заработано"]
        ws.append(headers)  # Заголовок таблицы

        for employee_name, employee_earnings in earnings_by_employee.items():
            row = [employee_name]
            for service_name in services:
                count = queryset.filter(personal__full_name=employee_name, service__title=service_name).count()
                row.append(count)
            row.append(f"{employee_earnings} руб.")
            ws.append(row)  # Строка с данными по сотруднику

        total_row = ["ИТОГО"] + [total_services[service] for service in services] + [f"{total_earnings} руб."]
        ws.append(total_row)  # Итоговая строка

        # Стиль ячеек
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.font = Font(name='Times New Roman', size=12)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                cell.alignment = Alignment(wrapText=True)

        # Автоматическое определение ширины столбцов
        for col_index, col in enumerate(ws.columns):
            max_width = 0
            for row_index, cell in enumerate(col):
                if row_index > 0:
                    if isinstance(cell, openpyxl.cell.cell.Cell):
                        if cell.value is not None:
                            cell.font = Font(name='Times New Roman', size=12)
                            cell.alignment = Alignment(wrapText=True)
                            rendered_width = len(str(cell.value)) * 1.5
                            max_width = max(max_width, rendered_width)

            col_letter = openpyxl.utils.get_column_letter(col_index + 1)
            ws.column_dimensions[col_letter].width = max_width

        ws.append([])  # Добавляем пустую строку перед подписью
        signature_text = "Начальник отдела маркетинга_________П.С.Стасов"
        ws.merge_cells('A' + str(ws.max_row + 1) + ':C' + str(ws.max_row + 1))
        signature_cell = ws.cell(row=ws.max_row + 1, column=1, value=signature_text)
        signature_cell.font = Font(name='Times New Roman', size=14)
        signature_cell.alignment = Alignment(horizontal='left')

        today = datetime.now().strftime("%Y-%m-%d")
        filename = f"Отчет_{today}.xlsx"
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quote(filename)}'
        try:
            wb.save(response)
            return response
        except Exception as e:
            messages.error(request, f"Ошибка при генерации Excel файла: {e}")
            return redirect('admin:index')

    def export_to_docx(self, request, start_date, end_date):
        queryset = self.get_queryset(request)
        services, earnings_by_employee, total_services, total_earnings = self.process_data(queryset, start_date, end_date)

        document = Document()
        heading_text = f'Отчет по услугам за период с {start_date.strftime("%d-%m-%Y")} по {end_date.strftime("%d-%m-%Y")}'
        heading = document.add_heading(level=1).add_run(heading_text)
        heading.font.name = 'Times New Roman'
        heading.font.size = Pt(16)
        heading.bold = True  # Заголовок отчета

        table = document.add_table(rows=1, cols=len(services) + 2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Сотрудник'

        for idx, service in enumerate(services):
            hdr_cells[idx + 1].text = service
        hdr_cells[-1].text = 'Заработано'  # Заголовки таблицы

        for employee_name, employee_earnings in earnings_by_employee.items():
            row_cells = table.add_row().cells
            row_cells[0].text = employee_name
            for service_name in services:
                count = queryset.filter(personal__full_name=employee_name, service__title=service_name).count()
                row_cells[services.index(service_name) + 1].text = str(count)
            row_cells[-1].text = str( f"{employee_earnings} руб.")  # Строка с данными по сотруднику

        total_row_cells = table.add_row().cells
        total_row_cells[0].text = 'ИТОГО'
        for idx, service in enumerate(services):
            total_row_cells[idx + 1].text = str(total_services[service])
        total_row_cells[-1].text = str(f"{total_earnings} руб.")  # Итоговая строка

        # Применение стиля шрифта ко всем ячейкам таблицы
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)  # Согласованный размер шрифта

        document.add_paragraph()
        signature_paragraph = document.add_paragraph("Начальник отдела маркетинга_________П.С.Стасов")
        signature_paragraph.alignment = 2  # Выравнивание по правому краю
        for run in signature_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)  # размер подпись

        today = datetime.now().strftime("%Y-%m-%d")
        filename = f"Отчет_{today}.docx"
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quote(filename)}'
        try:
            document.save(response)
            return response
        except Exception as e:
            messages.error(request, f"Ошибка при генерации DOCX файла: {e}")
            return redirect('admin:index')

    # Регистрация модели Client в админке
@admin.register(Client)
class ClientAdmin(admin.ModelAdmin):
    form = ClientForm  
    list_display = ('full_name', 'birth_date', 'phone')  # Названия столбцов
    list_filter = ('birth_date',)  # Фильтрация по дате рождения
    search_fields = ('full_name', 'phone')  # Поиск по ФИО и телефону


# Регистрация модели Personal в админке
@admin.register(Personal)
class PersonalAdmin(admin.ModelAdmin):
    form = PersonalForm  
    list_display = ('display_full_name', 'post', 'date_of_employment', 'status_color')  # Названия столбцов
    list_filter = ('post', 'date_of_employment', 'status')  # Фильтрация по должности и дате приема
    search_fields = ('full_name', 'post')  # Поиск по ФИО и должности
    actions = ['Установить_статус_Уволен','Установить_статус_Работает']

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.filter()  

    def Установить_статус_Уволен(self, request, queryset):
        queryset.update(status='уволен')

    def Установить_статус_Работает(self, request, queryset):
        queryset.update(status='работает')
    
    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        form.base_fields['date_of_employment'].label = 'Дата приема на работу'  
        return form

    def status_color(self, obj):
        if obj.status == 'работает':
            return format_html('<span style="color: green;">{}</span>', obj.status)
        elif obj.status == 'уволен':
            return format_html('<span style="color: red;">{}</span>', obj.status)  
    status_color.short_description = 'Статус'

    def display_full_name(self, obj):
        if obj.status == 'уволен':
            return format_html('<span style="color: red;">{}</span>', obj.full_name)
        return obj.full_name
    display_full_name.short_description = 'ФИО'

    def calendar_link(self, obj):
        return format_html('<a href="{}">Календарь</a>', reverse('work_schedule', args=[obj.id])) 
    calendar_link.short_description = 'Календарь'

    list_display = ('full_name', 'post', 'date_of_employment', 'status_color', 'calendar_link')

# Регистрация модели Service в админке
@admin.register(Service)
class ServiceAdmin(admin.ModelAdmin):
    list_display = ('title', 'duration', 'price', 'is_active')
    list_filter = ('duration', 'price','is_active')
    search_fields = ['title']
    actions = ['Отключить_услугу', 'Активировать_услугу']  # Добавляем действие смены активности
    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.filter()  # Показываем только активные услуги

    def Отключить_услугу(self, request, queryset):
        for service in queryset:
            service.is_active = False  # Устанавливаем статус неактивный
            service.save()

    def Активировать_услугу(self, request, queryset):
        for service in queryset:
            service.is_active = True  # Устанавливаем статус активный
            service.save()

    def delete(self, using=None, keep_parents=False):
        self.is_active = False  
        self.save()


    
# Регистрация модели Appointment в админке
@admin.register(Appointment)
class AppointmentAdmin(admin.ModelAdmin):
    form = AppointmentForm  
    list_display = ('client_full_name', 'personal_full_name', 'service', 'appointment_date', 'appointment_time', 'status_color')   
    list_filter = ('personal__full_name', 'appointment_date', 'status')  
    search_fields = ('client__full_name', 'personal__full_name')  
    actions = ['Установить_статус_Выполнено','Установить_статус_В_ожидании']

    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        form.base_fields['appointment_date'].label = 'Дата записи'  # Изменяем название поля
        form.base_fields['appointment_time'].label = 'Время записи'  # Изменяем название поля
        return form

    def client_full_name(self, obj):
        if obj.status == 'выполнено':
            return format_html('<span style="color: green;">{}</span>', obj.client.full_name)
        else:
            return format_html('<span style="color: blue;">{}</span>', obj.client.full_name)

    client_full_name.short_description = 'Клиент'

    def personal_full_name(self, obj):
        return obj.personal.full_name

    personal_full_name.short_description = 'Сотрудник'

    def status_color(self, obj):
        if obj.status == 'выполнено':
            return format_html('<span style="color: green;">{}</span>', obj.status)
        else:
            return format_html('<span style="color: blue;">{}</span>', obj.status)

    status_color.short_description = 'Статус'

    def Установить_статус_Выполнено(self, request, queryset):
        queryset.update(status='выполнено')

    def Установить_статус_В_ожидании(self, request, queryset):
        queryset.update(status='в_ожидании')

    def save_model(self, request, obj, form, change):
        try:
            obj.full_clean()
            super().save_model(request, obj, form, change)
        except ValidationError as e:
            messages.error(request, str(e))
            raise


